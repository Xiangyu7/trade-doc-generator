# -*- coding: utf-8 -*-
"""
核心逻辑模块（命令行脚本和网页后端共用）
  - generate_invoice : 套 CPH模板 生成 PI/CI/PL（xlsx + pdf）
  - merge_pdfs       : 把多个文件合并成一个清关 PDF
  - extract_po       : 尽量从 PO 里读出字段，预填表单（草稿，需人工核对）
"""
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path

import openpyxl
import pypdf
from docx import Document
from PIL import Image

APP_NAME = "TradeDocGenerator"
BASE = Path(__file__).resolve().parent
RESOURCE_BASE = Path(getattr(sys, "_MEIPASS", BASE))
RESOURCE_DIR = RESOURCE_BASE / "resources"
TEMPLATE_DIR = RESOURCE_DIR / "templates"
APP_DATA_DIR = Path(os.environ.get(
    "TRADE_DOC_DATA_DIR",
    Path(os.environ.get("LOCALAPPDATA", BASE)) / APP_NAME,
))
OUTDIR = APP_DATA_DIR / "输出单据"


def _resource_file(name: str) -> Path:
    candidates = [
        TEMPLATE_DIR / name,
        RESOURCE_DIR / name,
        BASE / name,
        BASE.parent / name,
    ]
    return next((p for p in candidates if p.exists()), candidates[0])


TEMPLATE = _resource_file("CPH模板.xlsx")

CREATE_NO_WINDOW = getattr(subprocess, "CREATE_NO_WINDOW", 0)
SUBPROCESS_KW = {"creationflags": CREATE_NO_WINDOW} if os.name == "nt" else {}


def _first_existing(paths):
    for p in paths:
        if p and Path(p).exists():
            return str(Path(p))
    return None


def _find_soffice() -> str | None:
    env_path = os.environ.get("SOFFICE_PATH") or os.environ.get("LIBREOFFICE_PATH")
    candidates = [
        env_path,
        RESOURCE_DIR / "LibreOfficePortable" / "App" / "libreoffice" / "program" / "soffice.exe",
        RESOURCE_DIR / "libreoffice" / "program" / "soffice.exe",
        BASE / "LibreOfficePortable" / "App" / "libreoffice" / "program" / "soffice.exe",
        BASE / "libreoffice" / "program" / "soffice.exe",
        Path(sys.executable).resolve().parent / "libreoffice" / "program" / "soffice.exe",
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
        Path("/usr/bin/libreoffice"),
        Path("/usr/local/bin/libreoffice"),
        Path("/opt/homebrew/bin/libreoffice"),
        shutil.which("soffice"),
        shutil.which("libreoffice"),
    ]
    return _first_existing(candidates)


SOFFICE = _find_soffice()

# 订单字段 -> 形式发票主表单元格（CI/PL 自动联动）
FIELD_MAP = {
    "发票号":         ("A4",  "Invoice No.:{}"),
    "发票日期":       ("C4",  "PI Date.:{}"),
    "ReferenceNo":   ("E4",  "Reference No.:{}"),
    "买方名称及地址":  ("D6",  "{}"),
    "通知方":         ("A14", "{}"),
    "品名及HS编码":   ("A20", "{}"),
    "数量KGS":        ("D20", "{}"),
    "单价USD":        ("E20", "{}"),
    "付款条款":       ("B25", "{}"),
    "贸易术语":       ("B26", "{}"),
    "包装方式":       ("B27", "{}"),
    "装货港":         ("B28", "{}"),
    "卸货港":         ("B29", "{}"),
    "唛头":           ("B30", "{}"),
}
PL_MAP = {"体积CBM": ("F20", "{}")}
NUMERIC = {"数量KGS", "单价USD", "体积CBM"}
INVOICE_SHEETS = ("形式发票", "发票B", "箱单 B")

# 订单表/表单用到的所有字段（顺序即展示顺序）
ORDER_FIELDS = list(FIELD_MAP.keys()) + list(PL_MAP.keys())


def safe_name(s: str) -> str:
    bad = '/\\:*?"<>|'
    return "".join("_" if ch in bad else ch for ch in str(s)).strip() or "无发票号"


def _fill(ws, order, field_map):
    for col, (cell, fmt) in field_map.items():
        val = order.get(col)
        if val in (None, ""):
            continue
        if col in NUMERIC:
            try:
                num = float(val)
                ws[cell] = int(num) if num.is_integer() else num
            except (TypeError, ValueError):
                ws[cell] = val
        else:
            ws[cell] = fmt.format(val)


def _set_print(ws):
    ws.print_area = "A1:F40"
    ws.page_setup.orientation = "portrait"
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr = openpyxl.worksheet.properties.PageSetupProperties(fitToPage=True)


def office_to_pdf(path: Path, outdir: Path) -> Path | None:
    """用 LibreOffice 把 xlsx/docx 等转成 pdf。"""
    path = Path(path)
    if not SOFFICE:
        return None
    suffix = path.suffix.lower()
    flag = "--calc" if suffix in (".xlsx", ".xls") else (
        "--writer" if suffix in (".docx", ".doc") else None)
    cmd = [SOFFICE, "--headless"]
    if flag:
        cmd.append(flag)
    cmd += ["--convert-to", "pdf", "--outdir", str(outdir), str(path)]
    try:
        subprocess.run(cmd, check=True, capture_output=True, timeout=120, **SUBPROCESS_KW)
        pdf = Path(outdir) / (path.stem + ".pdf")
        return pdf if pdf.exists() else None
    except Exception:
        return None


# 兼容旧名字
xlsx_to_pdf = office_to_pdf


def image_to_pdf(img_path: Path, outdir: Path) -> Path | None:
    """把 jpg/png 图片转成单页 pdf。"""
    try:
        out = Path(outdir) / (Path(img_path).stem + "_img.pdf")
        Image.open(img_path).convert("RGB").save(out, "PDF")
        return out if out.exists() else None
    except Exception:
        return None


def generate_invoice(order: dict, outdir: Path = OUTDIR):
    """生成一票订单的 PI/CI/PL。返回 (xlsx_path, pdf_path|None)。"""
    outdir.mkdir(exist_ok=True)
    wb = openpyxl.load_workbook(TEMPLATE)
    _fill(wb["形式发票"], order, FIELD_MAP)
    _fill(wb["箱单 B"], order, PL_MAP)
    for sn in INVOICE_SHEETS:
        _set_print(wb[sn])
    for sn in wb.sheetnames:
        if sn not in INVOICE_SHEETS:
            wb[sn].sheet_state = "hidden"
    wb.active = wb.sheetnames.index("形式发票")

    stem = safe_name(order.get("发票号") or "未命名")
    xlsx_out = outdir / f"{stem}_PI-CI-PL.xlsx"
    wb.save(xlsx_out)
    pdf_out = xlsx_to_pdf(xlsx_out, outdir)
    return xlsx_out, pdf_out


# 单据 key -> (展示名, 对应 sheet)。CI/PL 靠公式联动主表「形式发票」，
# 单独导出时把主表留在文件里（隐藏）即可让公式正常取值。
INVOICE_DOCS = {
    "PI": ("PI 形式发票", "形式发票"),
    "CI": ("CI 商业发票", "发票B"),
    "PL": ("PL 装箱单",   "箱单 B"),
}


def generate_invoice_docs(order: dict, which=None, outdir: Path = OUTDIR):
    """按需把 PI / CI / PL 各自导成独立文件（每个一份 xlsx + pdf）。
    which：要生成的单据 key 列表（PI/CI/PL 的子集），None 表示全部。
    返回 [{key, label, xlsx, pdf}]，pdf 为 None 表示转 PDF 失败。"""
    outdir.mkdir(exist_ok=True)
    keys = [k for k in INVOICE_DOCS if (which is None or k in which)]
    stem = safe_name(order.get("发票号") or "未命名")
    results = []
    for key in keys:
        label, target = INVOICE_DOCS[key]
        wb = openpyxl.load_workbook(TEMPLATE)
        _fill(wb["形式发票"], order, FIELD_MAP)   # 主表（CI/PL 公式引用它）
        _fill(wb["箱单 B"], order, PL_MAP)        # 箱单独立字段
        # 只给目标表设打印区域（LibreOffice 会打印所有设了打印区域的表），
        # 其余表隐藏且不设打印区域，这样导出的 PDF 只含目标这一张。
        _set_print(wb[target])
        for sn in wb.sheetnames:
            wb[sn].sheet_state = "visible" if sn == target else "hidden"
        wb.active = wb.sheetnames.index(target)
        xlsx_out = outdir / f"{stem}_{key}.xlsx"
        wb.save(xlsx_out)
        pdf_out = xlsx_to_pdf(xlsx_out, outdir)
        results.append({
            "key": key,
            "label": label,
            "xlsx": xlsx_out.name,
            "pdf": pdf_out.name if pdf_out else None,
        })
    return results


IMAGE_EXT = (".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp")


def merge_pdfs(file_paths, output_path: Path):
    """把多个文件（按给定顺序）合并成一个 PDF。
    PDF 直接合并；Excel/Word 用 LibreOffice 转；图片用 Pillow 转。
    返回 (输出路径, 被跳过的文件名列表)。"""
    writer = pypdf.PdfWriter()
    tmpdir = output_path.parent
    skipped = []
    for fp in file_paths:
        fp = Path(fp)
        suffix = fp.suffix.lower()
        if suffix == ".pdf":
            src = fp
        elif suffix in (".xlsx", ".xls", ".docx", ".doc"):
            src = office_to_pdf(fp, tmpdir)
        elif suffix in IMAGE_EXT:
            src = image_to_pdf(fp, tmpdir)
        else:
            src = None
        if src is None:
            skipped.append(fp.name)
            continue
        try:
            reader = pypdf.PdfReader(str(src))
            for page in reader.pages:
                writer.add_page(page)
        except Exception:
            skipped.append(fp.name)
    with open(output_path, "wb") as f:
        writer.write(f)
    return output_path, skipped


# ----------------- PO 自动识别（尽力而为，结果需人工核对） -----------------
def _pdf_text(path: Path) -> str:
    try:
        reader = pypdf.PdfReader(str(path))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    except Exception:
        return ""


# ---- OCR：扫描件 PDF / 图片。优先 RapidOCR，PDF 转图片用 PyMuPDF。----
def _pdf_to_images(path: Path, tmpdir: Path, dpi: int = 200) -> list[Path]:
    try:
        import fitz
    except Exception:
        return []
    zoom = dpi / 72
    matrix = fitz.Matrix(zoom, zoom)
    images = []
    try:
        doc = fitz.open(str(path))
        for idx, page in enumerate(doc, 1):
            pix = page.get_pixmap(matrix=matrix, alpha=False)
            out = tmpdir / f"p{idx:03d}.png"
            pix.save(str(out))
            images.append(out)
    except Exception:
        return []
    return images


def _rapidocr_image_text(img_path: Path) -> str:
    try:
        from rapidocr_onnxruntime import RapidOCR
    except Exception:
        try:
            from rapidocr import RapidOCR
        except Exception:
            return ""
    try:
        engine = _rapidocr_image_text._engine
    except AttributeError:
        try:
            engine = RapidOCR()
        except Exception:
            return ""
        _rapidocr_image_text._engine = engine
    try:
        result = engine(str(img_path))
    except Exception:
        return ""
    rows = result[0] if isinstance(result, tuple) else result
    texts = []
    for row in rows or []:
        if isinstance(row, (list, tuple)) and len(row) >= 2:
            texts.append(str(row[1]))
    return "\n".join(texts)


def _ocr_text(path: Path) -> str:
    """对扫描件 PDF 或图片做 OCR，返回识别出的文字。任何环节失败都返回 ''。"""
    path = Path(path)
    suffix = path.suffix.lower()
    import tempfile
    tmpdir = None
    try:
        if suffix == ".pdf":
            tmpdir = Path(tempfile.mkdtemp(prefix="po_ocr_"))
            imgs = _pdf_to_images(path, tmpdir)
        elif suffix in IMAGE_EXT:
            imgs = [path]
        else:
            return ""
        if not imgs:
            return ""
        return "\n".join(filter(None, (_rapidocr_image_text(i) for i in imgs)))
    except Exception:
        return ""
    finally:
        if tmpdir:
            shutil.rmtree(tmpdir, ignore_errors=True)


def _field_after(text: str, label_pat: str) -> str:
    """匹配 'Label: value' 或 'Label\\n: value'（值在同行或下一行），返回 value。"""
    m = re.search(label_pat + r"\s*:?\s*([^\n:][^\n]*)", text, re.I)
    return m.group(1).strip(" :：") if m else ""


# 数量/单价的单位词
_UNIT = r"(KGS?|KILOGRAM|METRIC\s*TON|MT|TON)"


def _is_ton(unit: str) -> bool:
    return bool(re.search(r"M\s*T|TON", unit, re.I)) and not re.search(r"KG", unit, re.I)


def extract_po(path: Path) -> dict:
    """从 PO 文件尽量提取字段，返回预填 dict（含 _scanned/_ocr/_text 元信息）。
    有文字层直接读；扫描件/图片回退到 OCR。结果是草稿，必须人工核对。"""
    path = Path(path)
    text = _pdf_text(path) if path.suffix.lower() == ".pdf" else ""
    ocr_used = False
    if not text.strip():
        text = _ocr_text(path)         # 扫描件/图片 -> OCR
        ocr_used = bool(text.strip())
    out = {"_scanned": not text.strip(), "_ocr": ocr_used, "_text": text}
    if not text.strip():
        return out  # OCR 也读不出，全靠手填

    U = text.upper()

    # 贸易术语 + 港口（允许 CFR (COST AND FREIGHT) ANTWERP 这种带括号说明）
    m = re.search(r'\b(EXW|FOB|FCA|CFR|CIF|CIP|CPT|DAP|DDP|DPU)\b[ \-]*(?:\([^)]*\)\s*)?([A-Za-z][A-Za-z ,]{1,30})', text, re.I)
    if m:
        incoterm = m.group(1).upper()
        port = m.group(2).strip(" ,-")
        out["贸易术语"] = f"{incoterm} {port}"
        if incoterm == "FOB":
            out["装货港"] = port
        elif incoterm in ("CIF", "CFR", "CIP", "CPT", "DAP", "DDP", "DPU"):
            out["卸货港"] = port

    # HS / HSN 编码
    hs = re.search(r'(?:HS\s*CODE|HSN/?SAC|HSN|HS)\D{0,8}(\d{4}[.\s]?\d{2}[.\s]?\d{0,4})', text, re.I)
    if not hs:
        hs = re.search(r'\b(\d{4}\.\d{2}\.\d{2,4})\b', text)  # 2916.14.0090 这种
    if not hs:
        hs = re.search(r'\b(2916\d{4})\b', text)  # 海关8位连写
    hs_code = hs.group(1) if hs else ""

    # PO / Reference / 采购合同号：遍历所有匹配，取第一个“含数字”的，
    # 避免把后面的标签词（如 Contract Date）误当成值。
    for mm in re.finditer(r'(?:PURCHASE ORDER NUMBER|PURCHASE CONTRACT NO|CONTRACT NO|PO\s*N[O0]\b|ORDER\s*N[O0]\b)\.?\s*:?\s*([A-Za-z0-9][A-Za-z0-9\-]{2,})', text, re.I):
        cand = mm.group(1).strip()
        if re.search(r'\d', cand):
            out["ReferenceNo"] = cand
            break

    # 付款条款：先按标签取值，取不到再扫关键词行
    pay = _field_after(text, r"(?:Payment\s*Terms?|付款条款?)")
    if not pay or len(pay) < 6:
        for line in text.splitlines():
            ls = line.strip().lstrip(":：").strip()
            if len(ls) < 8 or ls.endswith(":"):
                continue
            if re.search(r'(T/T|TT |ADVANCE|PREPAYMENT|PAYABLE|DAYS DUE|AGAINST (BL|DOCUMENT|SHIPP)|CASH AGAINST|L/C|LC )', ls, re.I) and len(ls) < 100:
                pay = ls
                break
    if pay:
        out["付款条款"] = pay.strip(" :：")

    # 包装方式
    pk = _field_after(text, r"(?:Packing|包装方式?)")
    if pk and len(pk) >= 3:
        out["包装方式"] = pk

    # 数量（带单位，吨 -> 公斤 ×1000）
    qty = re.search(r'(?:Quantity|数量)\D{0,12}?([\d.,]{1,15})\s*' + _UNIT, text, re.I)
    if not qty:
        qty = re.search(r'([\d.,]{2,15})\s*' + _UNIT + r'\b', text, re.I)
    if qty:
        val = _num(qty.group(1))
        if _is_ton(qty.group(2)) and isinstance(val, (int, float)):
            val = val * 1000
            val = int(val) if float(val).is_integer() else val
        out["数量KGS"] = val

    # 单价（带单位，每吨 -> 每公斤 ÷1000）
    pm = re.search(r'(?:Unit\s*Price|Price|单价)\D{0,12}?(?:USD|US\$|\$)?\s*([\d.,]{1,15})\s*(?:USD|US\$|\$)?\s*(?:/|per|每)?\s*' + _UNIT, text, re.I)
    if pm:
        pv = _num(pm.group(1))
        if _is_ton(pm.group(2)) and isinstance(pv, (int, float)):
            pv = round(pv / 1000.0, 6)
        out["单价USD"] = pv

    # 品名（+ 等级% + HS）。先抓完整化学名，抓不到再抓缩写
    prod = ""
    m = re.search(r'([0-9]?-?[A-Za-z][A-Za-z\- ]*?(?:METHACRYLATE|ACRYLATE)(?:\s+\d{1,3}\s*%)?)', text, re.I)
    if m:
        prod = re.sub(r'\s+', ' ', m.group(1)).strip()
    else:
        m = re.search(r'\b(\d?-?(?:HPMA|HEMA|HPA|HEA))\b', text, re.I)
        if m:
            prod = m.group(1).strip()
    if len(prod) > 60:
        prod = prod[:60].strip()
    if prod or hs_code:
        out["品名及HS编码"] = (prod + (f"\nHS CODE: {hs_code}" if hs_code else "")).strip()

    # —— AI 提取（MiniMax）：配了 .ai_key 就用，AI 抓到的字段覆盖上面的正则草稿；
    #    没配 key / 断网 / 调用失败 都会自动回退到正则结果，保证离线也能用。——
    try:
        import ai_extract
        if ai_extract.ai_available():
            ai_fields = ai_extract.ai_extract_po(text)
            if ai_fields:
                out["_ai"] = True
                out.update(ai_fields)   # AI 优先
    except Exception:
        pass

    return out


def _num(s):
    """把各种千分位/小数写法转成数字；判断不了就原样返回。
    支持 美式 26,400.00 / 欧式 2.440,00 / 欧式小数 24,00 / 1.55 / 20.000。"""
    s = str(s).strip()
    if not re.search(r"\d", s):
        return s
    if "," in s and "." in s:
        # 同时有逗号和点：靠后的那个是小数点，另一个是千分位
        if s.rfind(",") > s.rfind("."):
            s2 = s.replace(".", "").replace(",", ".")   # 欧式 2.440,00 -> 2440.00
        else:
            s2 = s.replace(",", "")                      # 美式 26,400.00 -> 26400.00
    elif "," in s:
        intpart, _, frac = s.rpartition(",")
        if "," not in intpart and len(frac) != 3:
            s2 = s.replace(",", ".")                     # 24,00 / 24,5 -> 小数
        else:
            s2 = s.replace(",", "")                      # 26,400 / 1,234,567 -> 千分位
    elif "." in s:
        if re.match(r'^\d{1,3}(\.\d{3})+$', s):
            s2 = s.replace(".", "")                      # 20.000 / 1.234.567 -> 欧式千分位
        else:
            s2 = s                                       # 1.55 -> 小数
    else:
        s2 = s
    try:
        f = float(s2)
        return int(f) if f.is_integer() else f
    except ValueError:
        return s


# ----------------- 危化品字段识别（扫描性能单/鉴定报告，尽力而为，需人工核对） -----------------
def extract_chem(paths) -> dict:
    """从 性能单 / 分类鉴定报告 等源文件尽量提取危化品字段，预填表单（草稿）。
    性能单一般有文字层；鉴定报告多为扫描件，回退 OCR。结果必须人工核对。
    返回 {字段: 值, _texts: [(文件名, 文字)], _ocr: bool}。"""
    texts, ocr_used = [], False
    for p in paths:
        p = Path(p)
        t = _pdf_text(p) if p.suffix.lower() == ".pdf" else ""
        if not t.strip():
            t = _ocr_text(p)
            if t.strip():
                ocr_used = True
        texts.append((p.name, t))
    full = "\n".join(t for _, t in texts)
    out = {}

    # 性能单号：编号XXXZ01（如 326N42130002248Z01）
    m = re.search(r"编号\s*([0-9A-Za-z]+Z\d+)", full) or re.search(r"\b(\d{3}[A-Z]\d{6,}Z\d+)\b", full)
    if m:
        out["性能单号"] = m.group(1)

    # UN 包装标记：1H1/Y1.5/150/25 … CN/Cxxxx … PI:xxx
    m = re.search(r"(1H1/Y[\d.]+/\d+/\d+)\s*\n?\s*(CN/\S+?)_?\s*\n?\s*(PI:\S+)", full)
    if m:
        out["UN标记"] = f"Un {m.group(1)}  {m.group(2)}  {m.group(3)}"

    # 分类鉴别报告编号：一长串数字（可带 -1 后缀）
    m = re.search(r"(?:No|编号|报告编号)[)\s:：]*(\d{12,}-?\d*)", full) or re.search(r"\b(\d{17}-?\d?)\b", full)
    if m:
        out["分类鉴别报告编号"] = m.group(1)

    # 联合国编号：UN 2922 这类四位
    m = re.search(r"(?:联合国编号|UN\s*编号|UN\s*No)[^0-9]{0,8}(\d{4})", full) or re.search(r"\b(2922)\b", full)
    if m:
        out["UN编号"] = m.group(1)

    out["_ocr"] = ocr_used
    out["_texts"] = texts
    return out


# ===================== 危化品文档生成 =====================
# —— 通用字段（符合性声明 / 责任声明 共用，与包装类型无关）——
# 默认值取自现有模板里的实际数据，用户按本票货修改
CHEM_DEFAULTS = {
    "品名全称":         "腐蚀性液体，毒性，未另作规定的。丙烯酸-2-羟基丙酯",
    "品名简称":         "丙烯酸-2-羟基丙酯",
    "出口国":           "巴西",
    "桶数":             "80",
    "总重量":           "16000",
    "HS编码":           "2916129090",
    "UN编号":           "2922",
    "分类鉴别报告编号":  "37000010252112663-1",
    "UN标记":           "Un 1H1/Y1.5/150/25  CN/C421222  PI:008",
    "合同号":           "DEL26051542",
    "批次号":           "2026052501",
    "签署日期":         "2026年05月27日",
}

# —— 厂检单字段：按包装类型各有一套默认值（切换包装时这些字段跟着变）——
CHEM_PKG_DEFAULTS = {
    "200kg": {
        "性能单号":   "326N42130002248Z01",
        "货物密度":   "1.044g/mL",
        "报验桶数":   "80",
        "抽样桶数":   "13",
        "报检吨数":   "16",
        "灌装日期":   "2026.05.25",
        "鉴定日期":   "2026年05月26日",
        "检验依据":   "《国际海运危险货物规则》及SN/T0370.3-2021《出口危险货物包装检验规程第3部分》",
        "吨桶包装标记": "",   # 200kg 用模板固定标记，无需填
    },
    "吨桶": {
        "性能单号":   "325N23070007765Z01",
        "货物密度":   "1.044g/mL",
        "报验桶数":   "17",
        "抽样桶数":   "5",
        "报检吨数":   "17",
        "灌装日期":   "2025.11.05",
        "鉴定日期":   "2025年11月05日",
        "检验依据":   "(《国际海运危险货物规则》及SN/T0987.2-2014《出口危险货物中型散装货物包装容器检验规程 第2部分》",
        "吨桶包装标记": "31HA1/Y/0925/CN/C231211/4038/2027/1060L/55kg/100kPa",
    },
}

# UN 标记里 Un...CN/...PI:... 间距各文档不一，用正则统一匹配
_UNMARK_RE = r"Un\s*1H1/Y1\.5/150/25\s+CN/\S+\s+PI:\S*"


def _chem_rules(template_name, d):
    """返回某模板的替换规则列表：(模式, 待匹配, 新值)。模式 'lit' 字面 / 're' 正则。
    锚点是模板里的现有示例值，replace 成本票数据。短数字靠单位后缀(桶/件/吨)唯一定位。"""
    P  = d.get("品名全称", "")
    PS = d.get("品名简称", "")
    if template_name == "出口危险化学品生产企业符合性声明.docx":
        return [
            ("re",  r"共\d+桶\d+千克", f"共{d.get('桶数','')}桶{d.get('总重量','')}千克"),
            ("re",  _UNMARK_RE,        d.get("UN标记")),
            ("lit", "2916129090",      d.get("HS编码")),
            ("lit", "37000010252112663-1", d.get("分类鉴别报告编号")),
            ("lit", "巴西",            d.get("出口国")),
            ("lit", "2922",            d.get("UN编号")),
            ("lit", "2026年05月27日",  d.get("签署日期")),
            ("lit", CHEM_DEFAULTS["品名全称"], P),
        ]
    if template_name == "产品质量安全责任自负声明.docx":
        return [
            ("lit", "DEL26051542",   d.get("合同号")),
            ("lit", "16000千克",     f"{d.get('总重量','')}千克"),
            ("lit", "2026052501",    d.get("批次号")),
            ("lit", "巴西",          d.get("出口国")),
            ("lit", "2026年05月27日", d.get("签署日期")),
            ("lit", CHEM_DEFAULTS["品名全称"], P),
        ]
    if template_name == "丙烯酸-2-羟基丙酯 厂检单吨桶.docx":
        return [
            ("lit", "1.044g/mL",                 d.get("货物密度")),
            ("lit", "31HA1/Y/0925/CN/C231211/4038/2027/1060L/55kg/100kPa", d.get("吨桶包装标记")),
            ("lit", "325N23070007765Z01",        d.get("性能单号")),
            ("lit", "(《国际海运危险货物规则》及SN/T0987.2-2014《出口危险货物中型散装货物包装容器检验规程 第2部分》", d.get("检验依据")),
            ("lit", "2025.11.05",                d.get("灌装日期")),
            ("lit", "2025年11月05日",            d.get("鉴定日期")),
            ("re",  r"报检数量共\d+件，\d+吨",   f"报检数量共{d.get('报验桶数','')}件，{d.get('报检吨数','')}吨"),
            ("lit", "17桶",                      f"{d.get('报验桶数','')}桶"),  # 报验数量
            ("lit", "5桶",                       f"{d.get('抽样桶数','')}桶"),  # 抽样数量
            ("lit", "丙烯酸-2-羟基丙酯",         PS),
        ]
    if template_name == "丙烯酸-2-羟基丙酯 厂检单-200kg桶装.docx":
        return [
            ("lit", "1.044g/mL",                 d.get("货物密度")),
            ("lit", "326N42130002248Z01",        d.get("性能单号")),
            ("lit", "《国际海运危险货物规则》及SN/T0370.3-2021《出口危险货物包装检验规程第3部分》", d.get("检验依据")),
            ("lit", "2026.05.25",                d.get("灌装日期")),
            ("lit", "2026年05月26日",            d.get("鉴定日期")),
            ("re",  r"报检数量共\d+，\d+吨",     f"报检数量共{d.get('报验桶数','')}，{d.get('报检吨数','')}吨"),
            ("lit", "80桶",                      f"{d.get('报验桶数','')}桶"),  # 报验数量
            ("lit", "13桶",                      f"{d.get('抽样桶数','')}桶"),  # 抽样数量
            ("lit", "丙烯酸-2-羟基丙酯",         PS),
        ]
    return []


# 始终生成的两个声明
CHEM_SHARED = [
    ("符合性声明",   "出口危险化学品生产企业符合性声明.docx"),
    ("责任自负声明", "产品质量安全责任自负声明.docx"),
]
# 厂检单：按包装类型二选一
CHEM_FACTORY = {
    "200kg": ("厂检单(200kg桶装)", "丙烯酸-2-羟基丙酯 厂检单-200kg桶装.docx"),
    "吨桶":  ("厂检单(吨桶)",      "丙烯酸-2-羟基丙酯 厂检单吨桶.docx"),
}


def _set_para(p, fix):
    full = "".join(r.text for r in p.runs)
    new = fix(full)
    if new != full and p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""


def _apply_rules(doc, rules):
    def fix(text):
        for mode, pat, val in rules:
            if val in (None, ""):
                continue
            if mode == "re":
                text = re.sub(pat, lambda m: str(val), text)
            else:
                text = text.replace(pat, str(val))
        return text
    for p in doc.paragraphs:
        _set_para(p, fix)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _set_para(p, fix)


def generate_chem_docs(data: dict, packaging: str = "200kg", outdir: Path = OUTDIR):
    """按本票货数据生成 符合性声明 + 责任声明 + 对应包装的厂检单（docx + pdf）。
    packaging：'200kg' 或 '吨桶'，决定用哪个厂检单模板。返回结果列表。"""
    outdir.mkdir(exist_ok=True)
    # 厂检单字段缺失的，用该包装类型的默认值兜底
    d = dict(CHEM_PKG_DEFAULTS.get(packaging, {}))
    d.update({k: v for k, v in data.items() if v not in (None, "")})
    tag = safe_name(d.get("合同号") or "危化品")

    docs = list(CHEM_SHARED)
    if packaging in CHEM_FACTORY:
        docs.append(CHEM_FACTORY[packaging])

    results = []
    for label, tpl in docs:
        src = _resource_file(tpl)
        if not src.exists():
            results.append({"label": label, "error": f"找不到模板 {tpl}"})
            continue
        doc = Document(str(src))
        _apply_rules(doc, _chem_rules(tpl, d))
        docx_out = outdir / f"{tag}_{label}.docx"
        doc.save(docx_out)
        pdf_out = office_to_pdf(docx_out, outdir)
        results.append({
            "label": label,
            "docx": docx_out.name,
            "pdf": pdf_out.name if pdf_out else None,
        })
    return results
